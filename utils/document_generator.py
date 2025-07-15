import os
import tempfile
import zipfile
from datetime import datetime
from typing import Dict, List, Any
import weasyprint
from docx import Document
from docx.shared import Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from jinja2 import Environment, FileSystemLoader, select_autoescape
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Generates documents in multiple formats from processed data"""
    
    def __init__(self):
        # Get the correct template directory path
        import os
        current_dir = os.path.dirname(os.path.abspath(__file__))
        project_root = os.path.dirname(current_dir)
        self.template_dir = os.path.join(project_root, "templates")
        self.temp_dir = tempfile.mkdtemp()
        
        # Initialize Jinja2 environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(self.template_dir),
            autoescape=select_autoescape(['html', 'xml'])
        )
        
        # Add custom filters
        self.jinja_env.filters['format_number'] = self._format_number
        self.jinja_env.filters['format_currency'] = self._format_currency
        self.jinja_env.filters['format_date'] = self._format_date
        
        # Import number_to_words from formatters
        from utils.formatters import number_to_words
        self.jinja_env.filters['number_to_words'] = number_to_words
    
    def generate_first_page(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Generate First Page Summary in all formats"""
        try:
            template = self.jinja_env.get_template('first_page.html')
            html_content = template.render(data=data)
            
            # Generate files
            files = self._generate_all_formats(html_content, 'first_page', orientation='portrait')
            
            logger.info("First page generated successfully")
            return files
            
        except Exception as e:
            logger.error(f"Error generating first page: {str(e)}")
            raise
    
    def generate_deviation_statement(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Generate Deviation Statement in all formats"""
        try:
            template = self.jinja_env.get_template('deviation_statement.html')
            html_content = template.render(data=data)
            
            # Generate files (landscape orientation)
            files = self._generate_all_formats(html_content, 'deviation_statement', orientation='landscape')
            
            logger.info("Deviation statement generated successfully")
            return files
            
        except Exception as e:
            logger.error(f"Error generating deviation statement: {str(e)}")
            raise
    
    def generate_note_sheet(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Generate Note Sheet in all formats"""
        try:
            template = self.jinja_env.get_template('note_sheet.html')
            html_content = template.render(data=data)
            
            # Generate files
            files = self._generate_all_formats(html_content, 'note_sheet', orientation='portrait')
            
            logger.info("Note sheet generated successfully")
            return files
            
        except Exception as e:
            logger.error(f"Error generating note sheet: {str(e)}")
            raise
    
    def generate_extra_items(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Generate Extra Items in all formats"""
        try:
            template = self.jinja_env.get_template('extra_items.html')
            html_content = template.render(data=data)
            
            # Generate files
            files = self._generate_all_formats(html_content, 'extra_items', orientation='portrait')
            
            logger.info("Extra items generated successfully")
            return files
            
        except Exception as e:
            logger.error(f"Error generating extra items: {str(e)}")
            raise
    
    def generate_certificate(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Generate Certificate in all formats"""
        try:
            template = self.jinja_env.get_template('certificate.html')
            html_content = template.render(data=data)
            
            # Generate files
            files = self._generate_all_formats(html_content, 'certificate', orientation='portrait')
            
            logger.info("Certificate generated successfully")
            return files
            
        except Exception as e:
            logger.error(f"Error generating certificate: {str(e)}")
            raise
    
    def generate_memorandum(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Generate Memorandum in all formats"""
        try:
            template = self.jinja_env.get_template('memorandum.html')
            html_content = template.render(data=data)
            
            # Generate files
            files = self._generate_all_formats(html_content, 'memorandum', orientation='portrait')
            
            logger.info("Memorandum generated successfully")
            return files
            
        except Exception as e:
            logger.error(f"Error generating memorandum: {str(e)}")
            raise
    
    def create_combined_documents(self, documents: Dict[str, Dict[str, str]], data: Dict[str, Any]) -> Dict[str, str]:
        """Create combined documents from individual documents"""
        try:
            # Order of documents for combination
            doc_order = ['first_page', 'deviation_statement', 'note_sheet', 'extra_items', 'certificate', 'memorandum']
            
            # Create combined PDF
            combined_pdf = self._create_combined_pdf(documents, doc_order)
            
            # Create combined Word document
            combined_docx = self._create_combined_docx(documents, doc_order)
            
            # Create combined HTML
            combined_html = self._create_combined_html(documents, doc_order)
            
            return {
                'combined_pdf': combined_pdf,
                'combined_docx': combined_docx,
                'combined_html': combined_html
            }
            
        except Exception as e:
            logger.error(f"Error creating combined documents: {str(e)}")
            raise
    
    def create_zip_package(self, combined_files: Dict[str, str], data: Dict[str, Any]) -> str:
        """Create ZIP package with all generated files"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            agreement_no = data.get('title_info', {}).get('agreement_no', 'N_A')
            
            # Ensure the temp directory exists
            os.makedirs(self.temp_dir, exist_ok=True)
            
            zip_filename = os.path.join(self.temp_dir, f"bill_package_{agreement_no}_{timestamp}.zip")
            
            with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
                # Add combined files
                for file_type, file_path in combined_files.items():
                    if os.path.exists(file_path):
                        zipf.write(file_path, f"combined_{file_type.split('_')[-1]}")
            
            logger.info(f"ZIP package created: {zip_filename}")
            return zip_filename
            
        except Exception as e:
            logger.error(f"Error creating ZIP package: {str(e)}")
            raise
    
    def _generate_all_formats(self, html_content: str, doc_name: str, orientation: str = 'portrait') -> Dict[str, str]:
        """Generate PDF, DOCX, and HTML files from HTML content"""
        files = {}
        
        # Save HTML file
        html_file = os.path.join(self.temp_dir, f"{doc_name}.html")
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        files['html'] = html_file
        
        # Generate PDF
        pdf_file = os.path.join(self.temp_dir, f"{doc_name}.pdf")
        try:
            # Configure CSS for page setup
            css_string = f"""
            @page {{
                size: A4 {orientation};
                margin: 10mm;
            }}
            body {{
                font-family: Arial, sans-serif;
                font-size: 12px;
                line-height: 1.4;
            }}
            """
            
            weasyprint.HTML(string=html_content).write_pdf(
                pdf_file,
                stylesheets=[weasyprint.CSS(string=css_string)]
            )
            files['pdf'] = pdf_file
            
        except Exception as e:
            logger.error(f"Error generating PDF for {doc_name}: {str(e)}")
        
        # Generate DOCX
        docx_file = os.path.join(self.temp_dir, f"{doc_name}.docx")
        try:
            self._html_to_docx(html_content, docx_file, orientation)
            files['docx'] = docx_file
            
        except Exception as e:
            logger.error(f"Error generating DOCX for {doc_name}: {str(e)}")
        
        return files
    
    def _html_to_docx(self, html_content: str, output_file: str, orientation: str = 'portrait'):
        """Convert HTML content to DOCX format"""
        try:
            doc = Document()
            
            # Set page orientation and margins
            section = doc.sections[0]
            section.page_height = Mm(297)  # A4 height
            section.page_width = Mm(210)   # A4 width
            
            if orientation == 'landscape':
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_height = Mm(210)
                section.page_width = Mm(297)
            
            # Set margins to 10mm
            section.top_margin = Mm(10)
            section.bottom_margin = Mm(10)
            section.left_margin = Mm(10)
            section.right_margin = Mm(10)
            
            # Simple HTML to DOCX conversion
            # This is a basic implementation - for production use, consider using python-docx-template
            lines = html_content.split('\n')
            for line in lines:
                if line.strip():
                    # Remove HTML tags for basic conversion
                    clean_line = self._strip_html_tags(line)
                    if clean_line:
                        p = doc.add_paragraph(clean_line)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            doc.save(output_file)
            
        except Exception as e:
            logger.error(f"Error converting HTML to DOCX: {str(e)}")
            raise
    
    def _strip_html_tags(self, html_text: str) -> str:
        """Strip HTML tags from text"""
        import re
        clean = re.compile('<.*?>')
        return re.sub(clean, '', html_text).strip()
    
    def _create_combined_pdf(self, documents: Dict[str, Dict[str, str]], doc_order: List[str]) -> str:
        """Create combined PDF from individual PDFs"""
        try:
            import PyPDF2
            
            combined_pdf_file = os.path.join(self.temp_dir, "combined_bill.pdf")
            
            with open(combined_pdf_file, 'wb') as output_file:
                pdf_writer = PyPDF2.PdfWriter()
                
                for doc_name in doc_order:
                    if doc_name in documents and 'pdf' in documents[doc_name]:
                        pdf_path = documents[doc_name]['pdf']
                        if os.path.exists(pdf_path):
                            with open(pdf_path, 'rb') as pdf_file:
                                pdf_reader = PyPDF2.PdfReader(pdf_file)
                                for page in pdf_reader.pages:
                                    pdf_writer.add_page(page)
                
                pdf_writer.write(output_file)
            
            return combined_pdf_file
            
        except Exception as e:
            logger.error(f"Error creating combined PDF: {str(e)}")
            # Fallback: return first available PDF
            for doc_name in doc_order:
                if doc_name in documents and 'pdf' in documents[doc_name]:
                    return documents[doc_name]['pdf']
            raise
    
    def _create_combined_docx(self, documents: Dict[str, Dict[str, str]], doc_order: List[str]) -> str:
        """Create combined DOCX from individual DOCX files"""
        try:
            from docx import Document
            from docx.oxml import OxmlElement
            
            combined_docx_file = os.path.join(self.temp_dir, "combined_bill.docx")
            combined_doc = Document()
            
            for doc_name in doc_order:
                if doc_name in documents and 'docx' in documents[doc_name]:
                    docx_path = documents[doc_name]['docx']
                    if os.path.exists(docx_path):
                        doc = Document(docx_path)
                        
                        # Add page break before each document (except first)
                        if doc_name != doc_order[0]:
                            combined_doc.add_page_break()
                        
                        # Copy paragraphs
                        for paragraph in doc.paragraphs:
                            new_paragraph = combined_doc.add_paragraph(paragraph.text)
                            new_paragraph.style = paragraph.style
                        
                        # Copy tables
                        for table in doc.tables:
                            new_table = combined_doc.add_table(
                                rows=len(table.rows),
                                cols=len(table.columns)
                            )
                            
                            for i, row in enumerate(table.rows):
                                for j, cell in enumerate(row.cells):
                                    new_table.cell(i, j).text = cell.text
            
            combined_doc.save(combined_docx_file)
            return combined_docx_file
            
        except Exception as e:
            logger.error(f"Error creating combined DOCX: {str(e)}")
            # Fallback: return first available DOCX
            for doc_name in doc_order:
                if doc_name in documents and 'docx' in documents[doc_name]:
                    return documents[doc_name]['docx']
            raise
    
    def _create_combined_html(self, documents: Dict[str, Dict[str, str]], doc_order: List[str]) -> str:
        """Create combined HTML from individual HTML files"""
        try:
            combined_html_file = os.path.join(self.temp_dir, "combined_bill.html")
            
            with open(combined_html_file, 'w', encoding='utf-8') as output_file:
                output_file.write("""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Combined Bill</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; }
        .page-break { page-break-before: always; }
        .document-section { margin-bottom: 40px; }
    </style>
</head>
<body>
                """)
                
                for i, doc_name in enumerate(doc_order):
                    if doc_name in documents and 'html' in documents[doc_name]:
                        html_path = documents[doc_name]['html']
                        if os.path.exists(html_path):
                            if i > 0:
                                output_file.write('<div class="page-break"></div>')
                            
                            output_file.write(f'<div class="document-section">')
                            
                            with open(html_path, 'r', encoding='utf-8') as input_file:
                                content = input_file.read()
                                # Extract body content
                                body_start = content.find('<body>')
                                body_end = content.find('</body>')
                                if body_start != -1 and body_end != -1:
                                    body_content = content[body_start + 6:body_end]
                                    output_file.write(body_content)
                                else:
                                    output_file.write(content)
                            
                            output_file.write('</div>')
                
                output_file.write("""
</body>
</html>
                """)
            
            return combined_html_file
            
        except Exception as e:
            logger.error(f"Error creating combined HTML: {str(e)}")
            raise
    
    def _format_number(self, value: Any, decimals: int = 2) -> str:
        """Format number with specified decimal places"""
        try:
            if isinstance(value, (int, float)):
                return f"{value:.{decimals}f}"
            elif isinstance(value, str):
                return f"{float(value):.{decimals}f}"
        except:
            pass
        return "0.00"
    
    def _format_currency(self, value: Any) -> str:
        """Format currency with Indian numbering system"""
        try:
            if isinstance(value, (int, float)):
                return f"₹{value:,.2f}"
            elif isinstance(value, str):
                return f"₹{float(value):,.2f}"
        except:
            pass
        return "₹0.00"
    
    def _format_date(self, date_str: str) -> str:
        """Format date to dd/mm/yyyy format"""
        try:
            if isinstance(date_str, str) and date_str not in ['N/A', '']:
                # Try to parse different date formats
                from datetime import datetime
                
                # Try different formats
                formats = [
                    '%Y-%m-%d',
                    '%d/%m/%Y',
                    '%d-%m-%Y',
                    '%Y/%m/%d'
                ]
                
                for fmt in formats:
                    try:
                        parsed_date = datetime.strptime(date_str, fmt)
                        return parsed_date.strftime('%d/%m/%Y')
                    except:
                        continue
        except:
            pass
        
        return date_str
